from docx import Document
import os

class DocFileHandler:
    def __init__(self, file_path):
        self.file_path = file_path

    def read_text(self):
        doc = Document(self.file_path)
        text_content = ""
        for paragraph in doc.paragraphs:
            text_content += paragraph.text
        return text_content

    def write_text(self, content):
        doc = Document()
        doc.add_paragraph(content)
        doc.save(self.file_path)

    def add_heading(self, heading):
        doc = Document(self.file_path)
        doc.add_heading(heading, level=1)
        doc.save(self.file_path)

    def add_table(self, data):
        doc = Document(self.file_path)
        table = doc.add_table(rows=1, cols=len(data[0]))
        for i, row in enumerate(data):
            for j, cell in enumerate(row):
                table.cell(i, j).text = cell
        doc.save(self.file_path)
